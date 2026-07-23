"""Generate PLATFORM_CREDENTIALS.docx Word document."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any


def generate_credentials_doc(creds: dict[str, Any], output_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches, Pt
    except ImportError as e:
        raise RuntimeError("python-docx required: pip install python-docx") from e

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    title = doc.add_heading("AI Procurement Platform — Credentials & Access", 0)
    title.alignment = 0
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("CONFIDENTIAL — Local development credentials. Do not share publicly.")

    # Platform URLs
    doc.add_heading("Platform URLs", level=1)
    plat = creds.get("platform", {})
    _add_kv_table(doc, [
        ("Platform", plat.get("name", "")),
        ("Version", plat.get("version", "")),
        ("Landing Page", plat.get("landing_url", "")),
        ("CRM Dashboard", plat.get("dashboard_url", "")),
        ("CRM Login", plat.get("crm_login_url", "")),
        ("Marketing Page", plat.get("marketing_url", "")),
        ("Database", plat.get("database", "")),
    ])

    # Platform admin
    doc.add_heading("Platform Super Admin", level=1)
    admin = creds.get("platform_admin", {})
    _add_kv_table(doc, [
        ("Email", admin.get("email", "")),
        ("Password", admin.get("password", "")),
        ("Role", admin.get("role", "")),
        ("Name", admin.get("name", "")),
    ])

    # Agency
    doc.add_heading("Primary Agency Tenant", level=1)
    agency = creds.get("agency", {})
    _add_kv_table(doc, [
        ("Agency Name", agency.get("name", "")),
        ("Tenant ID", agency.get("tenant_id", "")),
        ("Store Slug", agency.get("slug", "")),
        ("Store URL", agency.get("store_url", "")),
        ("Default Margin %", str(agency.get("margin_percent", ""))),
    ])

    # Tenant stores
    doc.add_heading("AI Product Finder Stores (Tenants)", level=1)
    tenants = creds.get("tenants", [])
    if tenants:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Tenant Name"
        hdr[1].text = "Slug"
        hdr[2].text = "Plan"
        hdr[3].text = "Margin %"
        hdr[4].text = "Store URL"
        for t in tenants:
            row = table.add_row().cells
            row[0].text = str(t.get("name", ""))
            row[1].text = str(t.get("slug", ""))
            row[2].text = str(t.get("plan", ""))
            row[3].text = str(t.get("margin_percent", ""))
            row[4].text = str(t.get("store_url", ""))
    else:
        doc.add_paragraph("No tenants seeded yet. Run: python -m src.main seed-demo")

    # Partner storefront catalog
    doc.add_heading("Partner Storefront Catalog", level=1)
    doc.add_paragraph(
        "Each tenant has a branded AI Product Finder store with featured products, "
        "product photos, and specialty categories."
    )
    store_tenants = [t for t in tenants if t.get("slug") not in (None, "default")]
    if store_tenants:
        for t in store_tenants:
            doc.add_heading(str(t.get("name", "")), level=2)
            _add_kv_table(doc, [
                ("Store URL", str(t.get("store_url", ""))),
                ("Plan", str(t.get("plan", ""))),
                ("Margin %", str(t.get("margin_percent", ""))),
                ("Focus", str(t.get("tagline", ""))),
                ("Specialties", ", ".join(t.get("specialties") or [])),
            ])
            products = t.get("featured_products") or []
            if products:
                doc.add_paragraph("Featured products:")
                for name in products:
                    doc.add_paragraph(str(name), style="List Bullet")
    else:
        doc.add_paragraph("Run seed-demo to populate partner stores.")

    # Tenant user logins
    doc.add_heading("Tenant Admin Logins", level=1)
    users = creds.get("tenant_users", [])
    if users:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Tenant"
        hdr[1].text = "Name"
        hdr[2].text = "Email"
        hdr[3].text = "Password"
        hdr[4].text = "Role"
        for u in users:
            row = table.add_row().cells
            row[0].text = str(u.get("tenant", ""))
            row[1].text = str(u.get("name", ""))
            row[2].text = str(u.get("email", ""))
            row[3].text = str(u.get("password", ""))
            row[4].text = str(u.get("role", ""))

    # LLM
    doc.add_heading("LLM Configuration", level=1)
    llm = creds.get("llm", {})
    _add_kv_table(doc, [
        ("Provider", llm.get("provider", "")),
        ("Base URL", llm.get("base_url", "")),
        ("Model", llm.get("model", "")),
        ("OpenAI API Key", llm.get("openai_api_key", "")),
    ])

    # Email
    doc.add_heading("Email / SMTP", level=1)
    email = creds.get("email", {})
    _add_kv_table(doc, [
        ("SMTP Host", email.get("smtp_host", "")),
        ("SMTP Port", str(email.get("smtp_port", ""))),
        ("SMTP User", email.get("smtp_user", "")),
        ("SMTP Password", email.get("smtp_pass", "")),
        ("From Address", email.get("smtp_from", "")),
        ("Dry Run Mode", str(email.get("dry_run", ""))),
    ])

    # API
    doc.add_heading("Key API Endpoints", level=1)
    api = creds.get("api_endpoints", {})
    for k, v in api.items():
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    # Notes
    doc.add_heading("Important Notes", level=1)
    for note in creds.get("notes", []):
        doc.add_paragraph(note, style="List Bullet")

    doc.add_page_break()
    doc.add_heading("Quick Start", level=1)
    doc.add_paragraph("1. Start server: python -m src.main dashboard")
    doc.add_paragraph("2. Seed demo data: python -m src.main seed-demo")
    doc.add_paragraph("3. Open landing page: http://127.0.0.1:8765/")
    doc.add_paragraph("4. Super admin login: http://127.0.0.1:8765/login")
    doc.add_paragraph("5. Open demo store: http://127.0.0.1:8765/store?tenant=demo")
    doc.add_paragraph("6. CRM dashboard: http://127.0.0.1:8765/app")

    doc.save(str(output_path))
    return output_path


def _add_kv_table(doc, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = str(v)
    doc.add_paragraph("")


def generate_credentials_markdown(creds: dict[str, Any], output_path: Path) -> Path:
    """Write a readable markdown credentials reference."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    plat = creds.get("platform", {})
    admin = creds.get("platform_admin", {})
    lines = [
        "# AI Procurement Platform — Credentials",
        "",
        f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
        "",
        "> CONFIDENTIAL — Local development only. Do not commit or share publicly.",
        "",
        "## Platform URLs",
        "",
        f"| URL | Link |",
        f"|-----|------|",
        f"| Landing | {plat.get('landing_url', '')} |",
        f"| CRM Login | {plat.get('crm_login_url', '')} |",
        f"| CRM Dashboard | {plat.get('dashboard_url', '')} |",
        f"| Marketing | {plat.get('marketing_url', '')} |",
        "",
        "## Super Admin",
        "",
        f"- **Email:** `{admin.get('email', '')}`",
        f"- **Password:** `{admin.get('password', '')}`",
        f"- **Role:** {admin.get('role', '')}",
        f"- **Access:** Platform SaaS management, all tenants, seed demo, credentials",
        "",
        "## Tenant Admin Logins",
        "",
        "| Tenant | Name | Email | Password | Role |",
        "|--------|------|-------|----------|------|",
    ]
    for u in creds.get("tenant_users", []):
        lines.append(
            f"| {u.get('tenant', '')} | {u.get('name', '')} | `{u.get('email', '')}` | `{u.get('password', '')}` | {u.get('role', '')} |"
        )
    lines.extend([
        "",
        "## Demo Stores",
        "",
    ])
    for t in creds.get("tenants", []):
        if t.get("slug") in (None, "default"):
            continue
        lines.append(f"- **{t.get('name')}** — {t.get('store_url', '')}")
        specs = ", ".join(t.get("specialties") or [])
        if specs:
            lines.append(f"  - Specialties: {specs}")
        products = t.get("featured_products") or []
        if products:
            lines.append(f"  - Featured: {', '.join(products)}")
    lines.extend([
        "",
        "## Quick Start",
        "",
        "```bash",
        "python -m src.main dashboard",
        "python -m src.main seed-demo",
        "python -m src.main credentials",
        "```",
        "",
    ])
    for note in creds.get("notes", []):
        lines.append(f"- {note}")
    output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return output_path
